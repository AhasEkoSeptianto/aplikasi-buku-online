from docx import Document
import os
from datetime import datetime

# class untuk menghandel semua Script yang berhubungan dengan .Docx files
class Doc:
    def __init__(self):
        self.doc = Document()

    def CreateNoted(self,namaMhs,nimMhs,semester,matakuliah,prodi,kelas,opsiName="Noted " + datetime.today().strftime('%d-%m-%Y') + ".docx"):
        self.doc.add_paragraph("Nama\t\t: {}".format(namaMhs))
        self.doc.add_paragraph("Nim\t\t: {}".format(nimMhs))
        self.doc.add_paragraph("Semester\t\t: {}".format(semester))
        self.doc.add_paragraph("Matakuliah\t\t: {}".format(matakuliah))
        self.doc.add_paragraph("Kelas\t\t: {}".format(kelas))
        self.doc.add_paragraph("Prodi\t\t: {}".format(prodi))
        
        self.doc.save(opsiName)
        

    def SaveIdent(self, namaMhs, nimMhs,prodi, kelas):
        self.doc.add_paragraph("Nama\t\t: {}".format(namaMhs))
        self.doc.add_paragraph("Nim\t\t: {}".format(nimMhs))
        self.doc.add_paragraph("Kelas\t\t: {}".format(kelas))
        self.doc.add_paragraph("Prodi\t\t: {}".format(prodi))
        
        self.doc.save("Identify.docx")

    def CheckIdent(self):
        doc     = Document('info/Identify.docx')
        name    = (doc.paragraphs[0].text).replace('Nama\t\t: ','')
        nim     = (doc.paragraphs[1].text).replace('Nim\t\t: ','')
        kelas   = (doc.paragraphs[2].text).replace('Kelas\t\t: ','')
        prodi   = (doc.paragraphs[3].text).replace('Prodi\t\t: ','')
        
        ident = {
            'Name' : name,
            'Nim' : nim,
            'Kelas' : kelas,
            'Prodi' : prodi,
        }
        return ident

    def getBooks(self,semester,matakuliah,catatan=os.listdir()):
        doc = Document(catatan)
        myNoted = [] # list for contains books noted
        for Noted in doc.paragraphs:
            myNoted.append(Noted.text)

        return myNoted        

    def Update(self,Noted,text):
        self.doc = Document()#membuat lembar kerja baru
        self.doc.add_paragraph(text) #mengisi lembar kerja dengan text dari user
        self.doc.save(Noted) 
        

        